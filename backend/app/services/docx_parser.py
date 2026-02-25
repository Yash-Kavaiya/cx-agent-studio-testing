"""Document parser service - extracts test content from .docx files."""

import io
from typing import List
from docx import Document as DocxDocument


class DocxParser:
    @staticmethod
    def extract_text(file_bytes: bytes) -> str:
        doc = DocxDocument(io.BytesIO(file_bytes))
        sections = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue
            if "Heading" in style:
                level = style.replace("Heading", "").strip()
                sections.append(
                    f"{'#' * int(level) if level.isdigit() else '#'} {text}"
                )
            else:
                sections.append(text)
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                sections.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")
        return "\n\n".join(sections)

    @staticmethod
    def extract_sections(file_bytes: bytes) -> List[dict]:
        doc = DocxDocument(io.BytesIO(file_bytes))
        sections = []
        current_section = {"heading": "Introduction", "content": []}
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue
            if "Heading" in style:
                if current_section["content"]:
                    sections.append(current_section)
                current_section = {"heading": text, "content": []}
            else:
                current_section["content"].append(text)
        if current_section["content"]:
            sections.append(current_section)
        return sections


def parse_docx(file_bytes: bytes) -> str:
    """Convenience function to parse docx and extract text."""
    return DocxParser.extract_text(file_bytes)
